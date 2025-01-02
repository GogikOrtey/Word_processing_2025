from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

def process_document(file_path):
    try:
        doc = Document(file_path)
        print(f"🟦 Документ '{file_path}' успешно загружен.")
        print(f"Количество параграфов в документе: {len(doc.paragraphs)}")
        return doc
    except Exception as e:
        print(f"🟧 Ошибка при загрузке документа: {e}")
        return None

def modify_paragraph_style(paragraph):
    if len(paragraph.runs) == 0:
        return  # Пропустить параграф, если в нем нет runs
    run = paragraph.runs[0]
    run.font.name = 'Times New Roman'
    run.font.bold = True
    run.font.size = Pt(16)
    paragraph.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER

    # Set spacing before and after
    pPr = paragraph._element.get_or_add_pPr()
    spacing = OxmlElement('w:spacing')
    spacing.set(qn('w:before'), '240')  # 12pt before
    spacing.set(qn('w:after'), '60')    # 3pt after
    pPr.append(spacing)


def main(file_path):
    doc = process_document(file_path)
    if doc is None:
        return

    processing = False
    for paragraph in doc.paragraphs:
        if '%%%%%' in paragraph.text:
            processing = not processing
            continue
        if processing and paragraph.text.startswith('№'):
            paragraph.style = doc.styles['Heading 1']
            modify_paragraph_style(paragraph)
    doc.save('modified_document_3.docx')

# Path to your document
file_path = '7_1 Сборник бабушкиных стихов.docx'
main(file_path)
print("🟩 Обработка успешно завершена!")
































